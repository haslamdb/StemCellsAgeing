#!/usr/bin/env python3
"""Generate a .docx report with FDR-adjusted q-values for B6 pathway comparisons."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Style adjustments
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('FDR-Adjusted Q-Values for Vitamin B6 Pathway Comparisons', level=1)

# Intro paragraph
doc.add_paragraph(
    'This document provides the FDR-adjusted q-values (Benjamini-Hochberg correction) '
    'for the two vitamin B6 biosynthesis pathways across the requested comparisons and methods. '
    'All q-values were computed genome-wide across all pathways tested in each analysis.'
)

# --- Section: Multiple Testing Correction ---
doc.add_heading('Multiple Testing Correction Method', level=2)
doc.add_paragraph(
    'All adjusted p-values use the Benjamini-Hochberg (BH) procedure for controlling '
    'the false discovery rate (FDR). The correction was applied genome-wide across all '
    'pathways tested in each analysis:'
)
bullets = doc.add_paragraph()
bullets.style = 'List Bullet'
bullets.text = 'DESeq2: BH correction applied automatically by the results() function (~55 pathways tested)'
b2 = doc.add_paragraph()
b2.style = 'List Bullet'
b2.text = 'ALDEx2 Welch (Y vs O): BH correction via we.eBH column (~250+ pathways tested)'
b3 = doc.add_paragraph()
b3.style = 'List Bullet'
b3.text = 'ALDEx2 GLM (DY vs DO, batch-corrected): BH correction via pval_BH column (~250+ pathways tested)'

# --- Section: Y vs O ---
doc.add_heading('Young (Y) vs Old (O)', level=2)

table1 = doc.add_table(rows=5, cols=4, style='Light Shading Accent 1')
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
headers = ['Method', 'Pathway', 'P-value', 'Q-value (BH FDR)']
for i, h in enumerate(headers):
    cell = table1.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Data rows
data_yo = [
    ('DESeq2', 'PYRIDOXSYN-PWY', '0.026', '0.455'),
    ('DESeq2', 'PWY0-845', '0.0011', '0.031*'),
    ('ALDEx2 Welch', 'PYRIDOXSYN-PWY', '0.017', '0.546'),
    ('ALDEx2 Welch', 'PWY0-845', '0.002', '0.218'),
]

for row_idx, (method, pathway, pval, qval) in enumerate(data_yo, start=1):
    table1.rows[row_idx].cells[0].text = method
    table1.rows[row_idx].cells[1].text = pathway
    table1.rows[row_idx].cells[2].text = pval
    table1.rows[row_idx].cells[3].text = qval

doc.add_paragraph()
p_note = doc.add_paragraph('* Significant after genome-wide FDR correction (q < 0.05)')
p_note.runs[0].italic = True
p_note.runs[0].font.size = Pt(10)

# --- Section: DY vs DO ---
doc.add_heading('Donor Young (DY) vs Donor Old (DO) — Batch-Corrected', level=2)

doc.add_paragraph(
    'Both DESeq2 and ALDEx2 GLM models included the batch factor (Experiment) '
    'in the design formula: ~ Experiment + Groups.'
)

table2 = doc.add_table(rows=5, cols=4, style='Light Shading Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(headers):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

data_dydo = [
    ('DESeq2 (batch-corrected)', 'PYRIDOXSYN-PWY', '0.051', '0.829'),
    ('DESeq2 (batch-corrected)', 'PWY0-845', '0.071', '0.885'),
    ('ALDEx2 GLM (batch-corrected)', 'PYRIDOXSYN-PWY', '0.0027', '0.522'),
    ('ALDEx2 GLM (batch-corrected)', 'PWY0-845', '0.0022', '0.548'),
]

for row_idx, (method, pathway, pval, qval) in enumerate(data_dydo, start=1):
    table2.rows[row_idx].cells[0].text = method
    table2.rows[row_idx].cells[1].text = pathway
    table2.rows[row_idx].cells[2].text = pval
    table2.rows[row_idx].cells[3].text = qval

doc.add_paragraph()
p_note2 = doc.add_paragraph('None of the DY vs DO comparisons reach significance after genome-wide FDR correction.')
p_note2.runs[0].italic = True
p_note2.runs[0].font.size = Pt(10)

# --- Section: Interpretation ---
doc.add_heading('Summary', level=2)

doc.add_paragraph(
    'After genome-wide Benjamini-Hochberg FDR correction:'
)

s1 = doc.add_paragraph()
s1.style = 'List Bullet'
s1.text = (
    'PWY0-845 in Y vs O by DESeq2 is the only comparison that remains significant '
    'after genome-wide correction (q = 0.031).'
)

s2 = doc.add_paragraph()
s2.style = 'List Bullet'
s2.text = (
    'The ALDEx2 results for Y vs O show nominally significant p-values but do not '
    'survive genome-wide FDR correction (q = 0.218 for PWY0-845).'
)

s3 = doc.add_paragraph()
s3.style = 'List Bullet'
s3.text = (
    'The DY vs DO comparisons show consistent direction and nominally significant '
    'p-values for ALDEx2 (p = 0.002\u20130.003), but the genome-wide corrected q-values '
    'are 0.52\u20130.55 due to the large number of pathways tested.'
)

s4 = doc.add_paragraph()
s4.style = 'List Bullet'
s4.text = (
    'The relatively large q-values for ALDEx2 reflect the larger number of pathways '
    'tested (~250+) compared to DESeq2 (~55), which increases the multiple testing burden.'
)

# --- Section: Attached files ---
doc.add_heading('Source Data Files', level=2)
doc.add_paragraph('The following CSV files contain the complete results for all pathways tested:')

files = [
    ('DESeq2_Pathways_Y_vs_O_Full.csv', 'DESeq2 full results for Y vs O (padj column)'),
    ('ALDEx2_Pathways_Y_vs_O_Full.csv', 'ALDEx2 full results for Y vs O (we.eBH column)'),
    ('DESeq2_Pathways_DY_vs_DO_BatchCorrected_Full.csv', 'DESeq2 batch-corrected results for DY vs DO (padj column)'),
    ('ALDEx2_Pathways_DY_vs_DO_BatchCorrected_Full.csv', 'ALDEx2 GLM batch-corrected results for DY vs DO (pval_BH column)'),
]

for fname, desc in files:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    run = p.add_run(fname)
    run.bold = True
    p.add_run(f' — {desc}')

# Save
outpath = os.path.join(os.path.dirname(__file__), 'B6_Pathway_FDR_QValues.docx')
doc.save(outpath)
print(f'Saved to: {outpath}')
