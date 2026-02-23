#!/usr/bin/env python3
"""Generate ALDEx2 B6 pathway tables and .docx report for collaborator.

Both Y vs O and DY vs DO use ALDEx2 GLM with batch correction
(~ Experiment + Groups), providing consistent Estimate/SE/pval columns.
"""

import csv
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))

# ── Read ALDEx2 GLM Y vs O (new, with SE) ──
yo_glm_file = os.path.join(BASE, "response_to_reviewers", "ALDEx2_B6_Pathways_Y_vs_O_GLM.csv")
yo_b6 = []
with open(yo_glm_file) as f:
    reader = csv.DictReader(f)
    for row in reader:
        yo_b6.append(row)

# ── Read ALDEx2 GLM DY vs DO (batch-corrected, all columns) ──
dydo_file = os.path.join(BASE, "ALDEx2_GLM_DY_vs_DO_BatchCorrected_AllColumns.csv")
dydo_b6 = []
with open(dydo_file) as f:
    reader = csv.DictReader(f)
    for row in reader:
        pw = row["Pathway"]
        if "PYRIDOXSYN" in pw or "PWY0-845" in pw:
            dydo_b6.append(row)

# ── Update the DY vs DO CSV to match Y vs O format ──
dydo_csv_out = os.path.join(BASE, "response_to_reviewers", "ALDEx2_B6_Pathways_DY_vs_DO_GLM_BatchCorrected.csv")
with open(dydo_csv_out, "w", newline="") as f:
    writer = csv.writer(f)
    writer.writerow(["Pathway", "Estimate_CLR", "SE", "t.val", "pval", "pval_BH"])
    for row in dydo_b6:
        writer.writerow([
            row["Pathway"],
            row["GroupsDY:Est"], row["GroupsDY:SE"], row["GroupsDY:t.val"],
            row["GroupsDY:pval"], row["GroupsDY:pval.padj"]
        ])
print(f"Wrote: {dydo_csv_out}")

# ── Generate .docx ──
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

doc.add_heading('ALDEx2 GLM Results for Vitamin B6 Pathways', level=1)
doc.add_heading('For Supplementary Figure: Log2 Fold Change Plots', level=2)

# Intro
doc.add_paragraph(
    'Below are the ALDEx2 GLM results for the two vitamin B6 biosynthesis pathways '
    '(PYRIDOXSYN-PWY and PWY0-845) for both the Y vs O and DY vs DO comparisons. '
    'Both analyses use ALDEx2 GLM with batch correction (model: ~ Experiment + Groups), '
    'providing consistent output columns (Estimate, SE, p-value) suitable for generating '
    'log2FC bar plots with error bars.'
)

# Method section
doc.add_heading('Method', level=2)

doc.add_paragraph(
    'ALDEx2 performs centered log-ratio (CLR) transformation on the count data using '
    '128 Monte Carlo Dirichlet instances, then fits a generalized linear model (GLM) '
    'to each pathway. The model formula is:'
)
p_formula = doc.add_paragraph('~ Experiment + Groups')
p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p_formula.runs:
    run.bold = True
    run.font.size = Pt(12)

doc.add_paragraph(
    'where Experiment accounts for batch effects and Groups tests the biological comparison '
    '(Y vs O or DY vs DO). The reference level is O (for Y vs O) or DO (for DY vs DO), '
    'so positive estimates indicate higher abundance in Young/Donor-Young.'
)

# Column explanation
doc.add_heading('Column Descriptions', level=2)

col_table = doc.add_table(rows=6, cols=2, style='Light Shading Accent 1')
col_table.alignment = WD_TABLE_ALIGNMENT.CENTER
col_headers = ['Column', 'Description']
for i, h in enumerate(col_headers):
    cell = col_table.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True

col_data = [
    ('Estimate_CLR', 'GLM coefficient for the group effect, on the CLR (log2) scale. '
     'Equivalent to log2 fold change. Positive = higher in Y (or DY).'),
    ('SE', 'Standard error of the GLM coefficient. Use for error bars on bar plots.'),
    ('t.val', 'T-statistic (Estimate / SE).'),
    ('pval', 'Nominal p-value from the GLM.'),
    ('pval_BH', 'Benjamini-Hochberg FDR-adjusted p-value (genome-wide across all pathways).'),
]
for r, (col, desc) in enumerate(col_data, 1):
    col_table.rows[r].cells[0].text = col
    for run in col_table.rows[r].cells[0].paragraphs[0].runs:
        run.bold = True
    col_table.rows[r].cells[1].text = desc

# ── Table 1: Y vs O ──
doc.add_paragraph()
doc.add_heading('Table 1: ALDEx2 GLM — Young (Y) vs Old (O), Batch-Corrected', level=2)
doc.add_paragraph(
    'Model: ~ Experiment + Groups (reference = O). '
    'Y and O samples span batches 1 and 3 (6+5 per group). '
    'Both pathways are nominally significant.'
)

t1 = doc.add_table(rows=3, cols=6, style='Light Shading Accent 1')
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
t1_headers = ['Pathway', 'Estimate\n(log2FC)', 'SE', 't-value', 'p-value', 'q-value\n(BH FDR)']
for i, h in enumerate(t1_headers):
    t1.rows[0].cells[i].text = h
    for run in t1.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True

for r, row in enumerate(yo_b6, 1):
    if "PYRIDOXSYN" in row["Pathway"]:
        name = "PYRIDOXSYN-PWY"
    else:
        name = "PWY0-845"
    t1.rows[r].cells[0].text = name
    t1.rows[r].cells[1].text = f'{float(row["Estimate_CLR"]):.3f}'
    t1.rows[r].cells[2].text = f'{float(row["SE"]):.3f}'
    t1.rows[r].cells[3].text = f'{float(row["t.val"]):.3f}'
    t1.rows[r].cells[4].text = f'{float(row["pval"]):.4f}'
    t1.rows[r].cells[5].text = f'{float(row["pval_BH"]):.3f}'

# ── Table 2: DY vs DO ──
doc.add_paragraph()
doc.add_heading('Table 2: ALDEx2 GLM — Donor Young (DY) vs Donor Old (DO), Batch-Corrected', level=2)
doc.add_paragraph(
    'Model: ~ Experiment + Groups (reference = DO). '
    'Both pathways are nominally significant.'
)

t2 = doc.add_table(rows=3, cols=6, style='Light Shading Accent 1')
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(t1_headers):  # same headers
    t2.rows[0].cells[i].text = h
    for run in t2.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True

for r, row in enumerate(dydo_b6, 1):
    if "PYRIDOXSYN" in row["Pathway"]:
        name = "PYRIDOXSYN-PWY"
    else:
        name = "PWY0-845"
    t2.rows[r].cells[0].text = name
    t2.rows[r].cells[1].text = f'{float(row["GroupsDY:Est"]):.3f}'
    t2.rows[r].cells[2].text = f'{float(row["GroupsDY:SE"]):.3f}'
    t2.rows[r].cells[3].text = f'{float(row["GroupsDY:t.val"]):.3f}'
    t2.rows[r].cells[4].text = f'{float(row["GroupsDY:pval"]):.4f}'
    t2.rows[r].cells[5].text = f'{float(row["GroupsDY:pval.padj"]):.3f}'

# ── Plotting guidance ──
doc.add_paragraph()
doc.add_heading('Notes for Plotting', level=2)

n1 = doc.add_paragraph()
n1.style = 'List Bullet'
n1.text = (
    'For log2FC bar plots: use Estimate_CLR as the bar height for both comparisons. '
    'Values are on the log2 scale (CLR-transformed) and directly comparable between '
    'the two comparisons.'
)

n2 = doc.add_paragraph()
n2.style = 'List Bullet'
n2.text = (
    'For error bars: use the SE column directly. These are proper GLM standard errors '
    'from both comparisons, providing a consistent representation.'
)

n3 = doc.add_paragraph()
n3.style = 'List Bullet'
n3.text = (
    'Significance: both pathways are nominally significant (p < 0.05) in both comparisons. '
    'Genome-wide FDR-corrected q-values do not reach 0.05, which is expected given '
    'the large number of pathways tested (~417). See the separate FDR report for details.'
)

n4 = doc.add_paragraph()
n4.style = 'List Bullet'
n4.text = (
    'Both models include batch correction (Experiment as covariate). The Y vs O comparison '
    'spans batches 1 and 3; the DY vs DO comparison spans batches 1, 2, and 3.'
)

# ── Source files ──
doc.add_heading('Attached Source Files', level=2)
doc.add_paragraph('The following CSV files contain the B6 pathway results:')

files = [
    ('ALDEx2_B6_Pathways_Y_vs_O_GLM.csv',
     'ALDEx2 GLM results for Y vs O, batch-corrected (B6 pathways, with SE)'),
    ('ALDEx2_B6_Pathways_DY_vs_DO_GLM_BatchCorrected.csv',
     'ALDEx2 GLM results for DY vs DO, batch-corrected (B6 pathways, with SE)'),
]
for fname, desc in files:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    run = p.add_run(fname)
    run.bold = True
    p.add_run(f' \u2014 {desc}')

doc.add_paragraph()
doc.add_paragraph('Full genome-wide results for all pathways (for reference):')

full_files = [
    ('ALDEx2_Pathways_Y_vs_O_GLM_Full.csv',
     'All 417 pathways, Y vs O GLM'),
    ('ALDEx2_Pathways_DY_vs_DO_BatchCorrected_Full.csv',
     'All pathways, DY vs DO GLM'),
]
for fname, desc in full_files:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    run = p.add_run(fname)
    run.bold = True
    p.add_run(f' \u2014 {desc}')

# Save
outpath = os.path.join(BASE, "response_to_reviewers", "ALDEx2_B6_Pathway_Tables_for_Plotting.docx")
doc.save(outpath)
print(f"Wrote: {outpath}")
